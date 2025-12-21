
"""
AI-Powered DOCX Paystub Generator v5.4 - COMPLETE FIXED VERSION
✓ FIXED: N/A values replaced with actual calculations
✓ FIXED: Text alignment preservation
✓ FIXED: YTD hours calculations
"""

import os
import sys
import json
import re
import io
from typing import Dict, List, Tuple, Optional, Any
from datetime import datetime, timedelta

import streamlit as st
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import google.generativeai as genai

st.set_page_config(
    page_title="Paystub Generator v5.4",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #555;
        text-align: center;
        margin-bottom: 2rem;
    }
    .log-container {
        background-color: #000000;
        border: 1px solid #333;
        border-radius: 0.5rem;
        padding: 1.5rem;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        max-height: 600px;
        overflow-y: auto;
        color: #00ff00;
        line-height: 1.8;
    }
    .stButton>button {
        width: 100%;
        background-color: #1f77b4;
        color: white;
        font-weight: bold;
        padding: 0.5rem;
        border-radius: 0.5rem;
    }
    .limit-warning {
        background-color: #fff3cd;
        border: 2px solid #ffc107;
        border-radius: 0.5rem;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

YTD_LIMITS = {
    2025: {
        'cpp_max': 4034.10,
        'ei_max': 1077.48,
        'cpp_max_income': 68500.00,
        'ei_max_income': 63400.00,
    },
    2026: {
        'cpp_max': 4230.45,
        'ei_max': 1123.07,
        'cpp_max_income': 74600.00,
        'ei_max_income': 68900.00,
    }
}

class PerfectCalculator:
    @staticmethod
    def calculate(yearly_income: float, period: int, hours: float = 80.0,
                 overtime_hours: float = 0.0, vacation_hours: float = 0.0,
                 stat_holidays: float = 0.0, year: int = 2025) -> Dict:
        
        limits = YTD_LIMITS.get(year, YTD_LIMITS[2025])
        
        biweekly = yearly_income / 26.0
        hourly_rate = biweekly / 80.0
        
        regular = hourly_rate * hours
        overtime_rate = hourly_rate * 1.5
        overtime_pay = overtime_rate * overtime_hours
        stat_holiday_pay = hourly_rate * stat_holidays
        
        gross_before_vacation = regular + overtime_pay + stat_holiday_pay
        vacation_pay = gross_before_vacation * 0.0125
        gross_current = gross_before_vacation + vacation_pay
        
        cpp_current = gross_current * 0.057
        ei_current = gross_current * 0.017
        tax_current = gross_current * 0.15
        
        regular_ytd = regular * period
        overtime_ytd = overtime_pay * period
        vacation_ytd = vacation_pay * period
        stat_holiday_ytd = stat_holiday_pay * period
        gross_ytd = gross_current * period
        
        cpp_ytd_calculated = (cpp_current * period) + 133.0
        ei_ytd_calculated = (ei_current * period) + 23.0
        
        cpp_ytd = min(cpp_ytd_calculated, limits['cpp_max'])
        ei_ytd = min(ei_ytd_calculated, limits['ei_max'])
        
        cpp_limit_hit = cpp_ytd_calculated >= limits['cpp_max']
        ei_limit_hit = ei_ytd_calculated >= limits['ei_max']
        
        if period >= 17:
            if cpp_limit_hit and cpp_ytd >= limits['cpp_max']:
                cpp_current_adjusted = 0.0
            else:
                cpp_current_adjusted = cpp_current
                
            if ei_limit_hit and ei_ytd >= limits['ei_max']:
                ei_current_adjusted = 0.0
            else:
                ei_current_adjusted = ei_current
        else:
            cpp_current_adjusted = cpp_current
            ei_current_adjusted = ei_current
        
        tax_ytd = tax_current * period
        deductions_ytd = cpp_ytd + ei_ytd + tax_ytd
        net_ytd = gross_ytd - deductions_ytd
        
        hours_ytd = hours * period
        overtime_hours_ytd = overtime_hours * period
        vacation_hours_ytd = vacation_hours * period
        stat_holidays_ytd = stat_holidays * period
        
        return {
            'hourly_rate': hourly_rate,
            'hours': hours,
            'overtime_hours': overtime_hours,
            'overtime_rate': overtime_rate,
            'vacation_hours': vacation_hours,
            'stat_holidays': stat_holidays,
            'period': period,
            'year': year,
            'regular_current': regular,
            'overtime_current': overtime_pay,
            'vacation_current': vacation_pay,
            'stat_holiday_current': stat_holiday_pay,
            'gross_current': gross_current,
            'cpp_current': cpp_current_adjusted,
            'ei_current': ei_current_adjusted,
            'tax_current': tax_current,
            'deductions_current': cpp_current_adjusted + ei_current_adjusted + tax_current,
            'net_current': gross_current - (cpp_current_adjusted + ei_current_adjusted + tax_current),
            'regular_ytd': regular_ytd,
            'overtime_ytd': overtime_ytd,
            'vacation_ytd': vacation_ytd,
            'stat_holiday_ytd': stat_holiday_ytd,
            'gross_ytd': gross_ytd,
            'cpp_ytd': cpp_ytd,
            'ei_ytd': ei_ytd,
            'tax_ytd': tax_ytd,
            'deductions_ytd': deductions_ytd,
            'net_ytd': net_ytd,
            'hours_ytd': hours_ytd,
            'overtime_hours_ytd': overtime_hours_ytd,
            'vacation_hours_ytd': vacation_hours_ytd,
            'stat_holidays_ytd': stat_holidays_ytd,
            'cpp_limit_hit': cpp_limit_hit,
            'ei_limit_hit': ei_limit_hit,
            'cpp_max': limits['cpp_max'],
            'ei_max': limits['ei_max'],
            'cpp_ytd_calculated': cpp_ytd_calculated,
            'ei_ytd_calculated': ei_ytd_calculated,
        }

class AdvancedExtractor:
    @staticmethod
    def extract_all(doc: Document) -> Dict:
        data = {
            'raw_paragraphs': [],
            'raw_tables': [],
            'all_values': [],
            'structured_data': {}
        }
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                data['raw_paragraphs'].append({'index': idx, 'text': text, 'type': 'paragraph'})
                numbers = re.findall(r'\d+[,.]?\d*\.?\d+', text)
                for num in numbers:
                    data['all_values'].append({
                        'value': num,
                        'context': text,
                        'location': f'P{idx}',
                        'type': 'paragraph'
                    })
        
        for table_idx, table in enumerate(doc.tables):
            table_data = {'index': table_idx, 'rows': [], 'cells_flat': []}
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    try:
                        cell_text = cell.text.strip()
                    except:
                        cell_text = ""
                    
                    cell_info = {
                        'row': row_idx,
                        'col': cell_idx,
                        'text': cell_text,
                        'location': f'T{table_idx}R{row_idx}C{cell_idx}'
                    }
                    
                    row_data.append(cell_info)
                    table_data['cells_flat'].append(cell_info)
                    
                    numbers = re.findall(r'\d+[,.]?\d*\.?\d+', cell_text)
                    for num in numbers:
                        data['all_values'].append({
                            'value': num,
                            'context': cell_text,
                            'location': f'T{table_idx}R{row_idx}C{cell_idx}',
                            'type': 'table_cell',
                            'cell': cell
                        })
                
                table_data['rows'].append(row_data)
            data['raw_tables'].append(table_data)
        
        structured = []
        for p in data['raw_paragraphs']:
            structured.append(f"[PARAGRAPH {p['index']}] {p['text']}")
        
        for t in data['raw_tables']:
            structured.append(f"\n[TABLE {t['index']}]")
            for row in t['rows']:
                row_text = " | ".join([f"{c['location']}: {c['text']}" for c in row if c['text']])
                if row_text:
                    structured.append(row_text)
        
        data['structured_data'] = "\n".join(structured)
        return data

class AlignmentPreserver:
    @staticmethod
    def get_paragraph_format(paragraph):
        try:
            return {
                'alignment': paragraph.alignment,
                'left_indent': paragraph.paragraph_format.left_indent,
                'right_indent': paragraph.paragraph_format.right_indent,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
            }
        except:
            return {'alignment': None, 'left_indent': None, 'right_indent': None, 'space_before': None, 'space_after': None}
    
    @staticmethod
    def apply_paragraph_format(paragraph, fmt):
        try:
            if fmt.get('alignment') is not None:
                paragraph.alignment = fmt['alignment']
            if fmt.get('left_indent') is not None:
                paragraph.paragraph_format.left_indent = fmt['left_indent']
            if fmt.get('right_indent') is not None:
                paragraph.paragraph_format.right_indent = fmt['right_indent']
            if fmt.get('space_before') is not None:
                paragraph.paragraph_format.space_before = fmt['space_before']
            if fmt.get('space_after') is not None:
                paragraph.paragraph_format.space_after = fmt['space_after']
        except:
            pass
    
    @staticmethod
    def get_run_format(run):
        try:
            return {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size,
                'font_name': run.font.name,
            }
        except:
            return {'bold': None, 'italic': None, 'underline': None, 'font_size': None, 'font_name': None}
    
    @staticmethod
    def apply_run_format(run, fmt):
        try:
            if fmt.get('bold') is not None:
                run.bold = fmt['bold']
            if fmt.get('italic') is not None:
                run.italic = fmt['italic']
            if fmt.get('underline') is not None:
                run.underline = fmt['underline']
            if fmt.get('font_size'):
                run.font.size = fmt['font_size']
            if fmt.get('font_name'):
                run.font.name = fmt['font_name']
        except:
            pass

class OptimizedAI:
    def __init__(self, api_key: str, add_log_callback=None):
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.0-flash-exp', generation_config={"temperature": 0.1, "top_p": 0.95, "top_k": 40})
        self.add_log = add_log_callback or (lambda x: None)
    
    def intelligent_mapping(self, extracted_data: Dict, calculations: Dict) -> List[Dict]:
        self.add_log("   🧠 AI: Enhanced value mapping (N/A prevention)...")
        
        limit_info = ""
        if calculations.get('cpp_limit_hit') or calculations.get('ei_limit_hit'):
            limit_info = f"""⚠️ YTD LIMITS: CPP ${calculations['cpp_max']:,.2f} {"(MAX)" if calculations['cpp_limit_hit'] else ""}, EI ${calculations['ei_max']:,.2f} {"(MAX)" if calculations['ei_limit_hit'] else ""}"""
        
#         prompt = f"""You are an EXPERT PAYSTUB ANALYZER.

# DOCUMENT:
# {extracted_data['structured_data']}

# VALUES:
# **YTD HOURS (CRITICAL - USE FOR YTD HOUR FIELDS):**
# - Hours YTD: {calculations['hours_ytd']:.1f}
# - Overtime Hours YTD: {calculations['overtime_hours_ytd']:.1f}
# - Stat Holiday Hours YTD: {calculations['stat_holidays_ytd']:.1f}

# **CURRENT:**
# - Rate: {calculations['hourly_rate']:.2f}, Hours: {calculations['hours']:.1f}, OT Hours: {calculations['overtime_hours']:.1f}
# - Regular: {calculations['regular_current']:.2f}, OT: {calculations['overtime_current']:.2f}, Vacation: {calculations['vacation_current']:.2f}
# - Gross: {calculations['gross_current']:.2f}, CPP: {calculations['cpp_current']:.2f}, EI: {calculations['ei_current']:.2f}, Tax: {calculations['tax_current']:.2f}
# - Net: {calculations['net_current']:.2f}

# **YTD:**
# - Regular: {calculations['regular_ytd']:.2f}, OT: {calculations['overtime_ytd']:.2f}, Vacation: {calculations['vacation_ytd']:.2f}
# - Gross: {calculations['gross_ytd']:.2f}, CPP: {calculations['cpp_ytd']:.2f}, EI: {calculations['ei_ytd']:.2f}, Tax: {calculations['tax_ytd']:.2f}
# - Net: {calculations['net_ytd']:.2f}
# {limit_info}

# RULES:
# 1. ❌ NEVER use "N/A" as new_value
# 2. ✅ Hours YTD → {calculations['hours_ytd']:.1f}
# 3. ✅ OT Hours YTD → {calculations['overtime_hours_ytd']:.1f}
# 4. ✅ Match exact numbers with calculations
# 5. ✅ 2 decimals for money, 1 for hours

# Return JSON array:
# [{{"field_name": "Hours YTD", "old_value": "1528.00", "new_value": "{calculations['hours_ytd']:.1f}", "location": "T0R5C2", "confidence": "high"}}]"""


     
        prompt = f"""You are an EXPERT PAYSTUB DOCUMENT ANALYZER.

DOCUMENT CONTENT:
{extracted_data['structured_data']}

YOUR TASK: Find and map ALL numeric values to new calculated values.

NEW CALCULATED VALUES:
**RATES & HOURS:**
- Hourly Rate: {calculations['hourly_rate']:.2f}
- Regular Hours: {calculations['hours']:.1f}
- Overtime Hours: {calculations['overtime_hours']:.1f}
- Overtime Rate (1.5x): {calculations['overtime_rate']:.2f}
- Vacation Hours: {calculations['vacation_hours']:.1f}
- Stat Holiday Hours: {calculations['stat_holidays']:.1f}

**CURRENT PERIOD:**
- Regular Pay: {calculations['regular_current']:.2f}
- Overtime Pay: {calculations['overtime_current']:.2f}
- Vacation Pay (1.25%): {calculations['vacation_current']:.2f}
- Stat Holiday Pay: {calculations['stat_holiday_current']:.2f}
- Gross Pay: {calculations['gross_current']:.2f}
- CPP (5.7%): {calculations['cpp_current']:.2f}
- EI (1.7%): {calculations['ei_current']:.2f}
- Federal Tax (15%): {calculations['tax_current']:.2f}
- Total Deductions: {calculations['deductions_current']:.2f}
- Net Pay: {calculations['net_current']:.2f}

**YEAR TO DATE:**
- Regular Pay YTD: {calculations['regular_ytd']:.2f}
- Overtime Pay YTD: {calculations['overtime_ytd']:.2f}
- Vacation Pay YTD: {calculations['vacation_ytd']:.2f}
- Stat Holiday YTD: {calculations['stat_holiday_ytd']:.2f}
- Gross Pay YTD: {calculations['gross_ytd']:.2f}
- CPP YTD: {calculations['cpp_ytd']:.2f}
- EI YTD: {calculations['ei_ytd']:.2f}
- Federal Tax YTD: {calculations['tax_ytd']:.2f}
- Total Deductions YTD: {calculations['deductions_ytd']:.2f}
- Net Pay YTD: {calculations['net_ytd']:.2f}

{limit_info}

**INSTRUCTIONS:**
1. Find EVERY numeric value in the document
2. Identify what each value represents (rate, hours, pay, deduction, etc.)
3. Match it to the correct NEW calculated value
4. Preserve number formatting (commas, decimals)
5. Map "N/A" or "0.00" values if they should have actual numbers now

**IMPORTANT FIELD TYPES TO FIND:**
- Hourly/Rate values
- Hours (regular, overtime, vacation, stat holiday)
- Pay amounts (current period)
- YTD (year to date) amounts
- Deductions (CPP, EI, Tax)
- Gross and Net totals

Return ONLY a JSON array like this:
[
  {{
    "field_name": "Hourly Rate",
    "old_value": "140.58",
    "new_value": "{calculations['hourly_rate']:.2f}",
    "location": "T0R2C3",
    "confidence": "high"
  }},
  {{
    "field_name": "Regular Hours",
    "old_value": "80.00",
    "new_value": "{calculations['hours']:.1f}",
    "location": "T1R1C1",
    "confidence": "high"
  }}
]

Return ONLY the JSON array, nothing else:"""


        try:
            response = self.model.generate_content(prompt)
            json_text = self._clean_json(response.text)
            mappings = json.loads(json_text)
            
            if not isinstance(mappings, list):
                mappings = []
            
            clean_mappings = []
            rejected = 0
            
            for m in mappings:
                if not isinstance(m, dict) or not m.get('old_value') or not m.get('new_value'):
                    continue
                
                new_val = str(m.get('new_value', '')).strip().upper()
                if 'N/A' in new_val or new_val == 'NA':
                    self.add_log(f"      ❌ Rejected N/A: {m.get('field_name')}")
                    rejected += 1
                    continue
                
                clean_mappings.append(m)
            
            if rejected > 0:
                self.add_log(f"      ⚠️  Filtered {rejected} N/A mappings")
            
            mappings = clean_mappings
            
        except Exception as e:
            self.add_log(f"      ⚠️  AI error: {e}")
            mappings = []
        
        self.add_log(f"      ✅ {len(mappings)} valid mappings")
        return mappings
    
    def _clean_json(self, text: str) -> str:
        text = text.strip()
        text = re.sub(r'^```json\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'^```\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'```\s*$', '', text, flags=re.MULTILINE)
        match = re.search(r'\[.*\]', text, re.DOTALL)
        return match.group(0) if match else text

class SafeReplacer:
    @staticmethod
    def find_and_replace_safe(doc: Document, mappings: List[Dict], extracted_data: Dict, add_log_callback=None) -> int:
        add_log = add_log_callback or (lambda x: None)
        replaced = 0
        
        for mapping in mappings:
            try:
                old_val = str(mapping.get('old_value', '')).strip()
                new_val = str(mapping.get('new_value', '')).strip()
                field = mapping.get('field_name', 'Unknown')
                location = mapping.get('location', '')
                
                if not old_val or not new_val or new_val.upper() in ['N/A', 'NA']:
                    continue
                
                if location:
                    success = SafeReplacer._replace_by_location(doc, location, old_val, new_val)
                    if success:
                        add_log(f"      ✅ {field}: '{old_val}' → '{new_val}'")
                        replaced += 1
                        continue
                
                for val_info in extracted_data['all_values']:
                    if SafeReplacer._values_match(val_info['value'], old_val):
                        if val_info['type'] == 'table_cell' and 'cell' in val_info:
                            success = SafeReplacer._replace_in_cell_safe(val_info['cell'], old_val, new_val)
                            if success:
                                add_log(f"      ✅ {field}: '{old_val}' → '{new_val}'")
                                replaced += 1
                                break
            except Exception as e:
                add_log(f"      ⚠️  Skip {field}: {e}")
        
        return replaced
    
    @staticmethod
    def _replace_by_location(doc: Document, location: str, old_val: str, new_val: str) -> bool:
        match = re.match(r'T(\d+)R(\d+)C(\d+)', location)
        if not match:
            return False
        
        try:
            t_idx, r_idx, c_idx = map(int, match.groups())
            if t_idx >= len(doc.tables):
                return False
            table = doc.tables[t_idx]
            if r_idx >= len(table.rows):
                return False
            row = table.rows[r_idx]
            if c_idx >= len(row.cells):
                return False
            return SafeReplacer._replace_in_cell_safe(row.cells[c_idx], old_val, new_val)
        except:
            return False
    
    @staticmethod
    def _replace_in_cell_safe(cell: _Cell, old_text: str, new_text: str) -> bool:
        try:
            replaced = False
            for para in cell.paragraphs:
                para_format = AlignmentPreserver.get_paragraph_format(para)
                
                for run in para.runs:
                    run_format = AlignmentPreserver.get_run_format(run)
                    
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        AlignmentPreserver.apply_run_format(run, run_format)
                        replaced = True
                    else:
                        old_clean = re.sub(r'[^\d.]', '', old_text)
                        run_clean = re.sub(r'[^\d.]', '', run.text)
                        if old_clean and run_clean and old_clean == run_clean:
                            run.text = run.text.replace(old_text, new_text)
                            AlignmentPreserver.apply_run_format(run, run_format)
                            replaced = True
                
                if replaced:
                    AlignmentPreserver.apply_paragraph_format(para, para_format)
            return replaced
        except:
            return False
    
    @staticmethod
    def _values_match(val1: str, val2: str) -> bool:
        try:
            clean1 = re.sub(r'[^\d.]', '', str(val1))
            clean2 = re.sub(r'[^\d.]', '', str(val2))
            return clean1 and clean2 and clean1 == clean2
        except:
            return False

class Validator:
    @staticmethod
    def validate_all(calculations: Dict, mappings: List[Dict], doc: Document) -> Tuple[bool, List[str]]:
        errors = []
        
        if abs(calculations['gross_current'] - (calculations['regular_current'] + calculations['overtime_current'] + calculations['vacation_current'] + calculations['stat_holiday_current'])) > 0.02:
            errors.append("Gross calculation error")
        
        if not mappings:
            errors.append("No mappings created")
        
        na_count = 0
        for mapping in mappings:
            new_val = str(mapping.get('new_value', '')).upper()
            if 'N/A' in new_val or new_val == 'NA':
                errors.append(f"N/A in: {mapping.get('field_name')}")
                na_count += 1
        
        if na_count > 0:
            errors.append(f"{na_count} N/A mappings found")
        
        return len(errors) == 0, errors

class OptimizedSystem:
    def __init__(self, api_key: str, add_log_callback=None):
        self.calculator = PerfectCalculator()
        self.extractor = AdvancedExtractor()
        self.ai = OptimizedAI(api_key, add_log_callback)
        self.replacer = SafeReplacer()
        self.validator = Validator()
        self.add_log = add_log_callback or (lambda x: None)
    
    def generate(self, input_doc: Document, yearly_income: float, period: int, hours: float = 80.0, 
                 overtime_hours: float = 0.0, vacation_hours: float = 0.0, stat_holidays: float = 0.0, 
                 year: int = 2025) -> Tuple[Document, int, Dict]:
        
        self.add_log("="*90)
        self.add_log(f"PAYSTUB GENERATOR v5.4 FIXED ({year})")
        self.add_log("   ✅ N/A prevention • ✅ Alignment preservation • ✅ YTD hours")
        self.add_log("="*90)
        self.add_log("")
        
        self.add_log("📊 STEP 1: Calculating")
        self.add_log("-" * 90)
        calcs = self.calculator.calculate(yearly_income, period, hours, overtime_hours, vacation_hours, stat_holidays, year)
        self.add_log(f"   Period: {period}/26, Gross: ${calcs['gross_current']:,.2f}, Net: ${calcs['net_current']:,.2f}")
        self.add_log("   ✅ Calculated")
        self.add_log("")
        
        self.add_log("📖 STEP 2: Extracting")
        self.add_log("-" * 90)
        extracted = self.extractor.extract_all(input_doc)
        self.add_log(f"   Values: {len(extracted['all_values'])}")
        self.add_log("   ✅ Extracted")
        self.add_log("")
        
        self.add_log("🧠 STEP 3: AI Mapping")
        self.add_log("-" * 90)
        mappings = self.ai.intelligent_mapping(extracted, calcs)
        self.add_log(f"   ✅ Mappings: {len(mappings)}")
        self.add_log("")
        
        self.add_log("🔍 STEP 4: Validation")
        self.add_log("-" * 90)
        valid, errors = self.validator.validate_all(calcs, mappings, input_doc)
        if errors:
            for err in errors:
                self.add_log(f"      ⚠️  {err}")
        else:
            self.add_log("   ✅ Validated")
        self.add_log("")
        
        self.add_log("✏️  STEP 5: Safe Replacement")
        self.add_log("-" * 90)
        replaced = self.replacer.find_and_replace_safe(input_doc, mappings, extracted, self.add_log)
        self.add_log("")
        self.add_log(f"   ✅ Replaced: {replaced}/{len(mappings)}")
        self.add_log("")
        
        self.add_log("="*90)
        self.add_log("✅ COMPLETE!")
        self.add_log("="*90)
        self.add_log("")
        
        return input_doc, replaced, calcs

def calculate_period(start_date: datetime, end_date: datetime) -> int:
    year_start = datetime(start_date.year, 1, 1)
    days = (end_date - year_start).days
    return min((days // 14) + 1, 26)



def main():
    if 'logs' not in st.session_state:
        st.session_state.logs = []
    if 'output_docx' not in st.session_state:
        st.session_state.output_docx = None
    if 'output_filename' not in st.session_state:
        st.session_state.output_filename = ""
    if 'api_key_validated' not in st.session_state:
        st.session_state.api_key_validated = False
    if 'api_key' not in st.session_state:
        st.session_state.api_key = ""
    if 'last_calcs' not in st.session_state:
        st.session_state.last_calcs = None
    
    st.markdown('<div class="main-header">🚀 Paystub Generator v5.3 FIXED</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">✅ Single AI Call • Safe Document Handling • No Corruption</div>', unsafe_allow_html=True)
    
    if not st.session_state.api_key_validated:
        st.markdown("---")
        st.markdown("### 🔑 Enter Your Gemini API Key")
        st.info("👉 Get your free API key from: https://aistudio.google.com/app/apikey")
        
        api_key_input = st.text_input(
            "Gemini API Key",
            type="password",
            placeholder="AIzaSy...",
            help="Enter your Google Gemini API key"
        )
        
        col_a, col_b = st.columns([1, 1])
        
        with col_a:
            if st.button("✅ Validate & Start", type="primary"):
                if api_key_input and len(api_key_input) > 20:
                    try:
                        genai.configure(api_key=api_key_input)
                        model = genai.GenerativeModel('gemini-2.0-flash-exp')
                        response = model.generate_content("test")
                        
                        st.session_state.api_key = api_key_input
                        st.session_state.api_key_validated = True
                        st.success("✅ API Key validated!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Invalid API Key: {str(e)}")
                else:
                    st.error("❌ Please enter a valid API key")
        
        with col_b:
            if st.button("🔄 Reset"):
                st.session_state.api_key_validated = False
                st.session_state.api_key = ""
                st.rerun()
        
        st.stop()
    
    st.success("🔑 API Key: Connected ✅")
    
    with st.sidebar:
        st.header("⚙️ Settings")
        
        year = st.selectbox(
            "Tax Year",
            options=[2025, 2026],
            index=0
        )
        
        limits = YTD_LIMITS[year]
        st.info(f"""
**{year} YTD Limits:**
- CPP Max: ${limits['cpp_max']:,.2f}
- EI Max: ${limits['ei_max']:,.2f}
        """)
        
        st.markdown("---")
        st.markdown("""
### 📖 Fixed Issues
✅ Single AI call (no rate limits)
✅ Safe document handling
✅ Robust JSON parsing
✅ No document corruption
✅ Enhanced value detection

### 📊 Calculation Rules
- **CPP**: 5.7%
- **EI**: 1.7%
- **Vacation**: 1.25%
- **Tax**: 15%
        """)
        
        st.markdown("---")
        
        if st.button("🔄 Reset All", type="secondary"):
            st.session_state.logs = []
            st.session_state.output_docx = None
            st.session_state.output_filename = ""
            st.session_state.api_key_validated = False
            st.session_state.api_key = ""
            st.session_state.last_calcs = None
            st.rerun()
    
    # Main Content
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload Original Paystub")
        uploaded_file = st.file_uploader(
            "Choose a DOCX file",
            type=['docx'],
            help="Upload the original paystub DOCX file"
        )
        
        if uploaded_file:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            st.info(f"📊 File size: {len(uploaded_file.getvalue()) / 1024:.2f} KB")
    
    with col2:
        st.header("💰 Income & Period Details")
        
        yearly_income = st.number_input(
            "Yearly Income ($)",
            min_value=0.0,
            value=79870.22,
            step=100.0,
            format="%.2f",
            help="Enter total yearly income"
        )
        
        st.markdown("**📅 Pay Period (Bi-Weekly)**")
        col2a, col2b = st.columns(2)
        
        with col2a:
            start_date = st.date_input(
                "Start Date",
                value=datetime(2025, 1, 5),
                help="Select pay period start date"
            )
        
        with col2b:
            end_date = st.date_input(
                "End Date",
                value=datetime(2025, 1, 18),
                help="Select pay period end date"
            )
        
        if start_date and end_date:
            period = calculate_period(datetime.combine(start_date, datetime.min.time()), 
                                     datetime.combine(end_date, datetime.min.time()))
            st.info(f"📊 Calculated Period: **{period}** of 26")
            
            if period >= 17:
                st.warning(f"⚠️ Period {period}: YTD limits may apply for high income earners")
    
    # Hours & Details
    st.markdown("---")
    st.header("⏰ Hours & Earnings Details")
    
    col3a, col3b, col3c, col3d = st.columns(4)
    
    with col3a:
        regular_hours = st.number_input(
            "Regular Hours",
            min_value=0.0,
            value=80.0,
            step=0.5,
            format="%.2f",
            help="Regular hours (default: 80)"
        )
    
    with col3b:
        overtime_hours = st.number_input(
            "Overtime Hours (1.5x)",
            min_value=0.0,
            value=0.0,
            step=0.5,
            format="%.2f",
            help="Overtime at 1.5x rate"
        )
    
    with col3c:
        vacation_hours = st.number_input(
            "Vacation Hours",
            min_value=0.0,
            value=0.0,
            step=0.5,
            format="%.2f",
            help="Vacation hours"
        )
    
    with col3d:
        stat_holiday_hours = st.number_input(
            "Stat Holiday Hours",
            min_value=0.0,
            value=0.0,
            step=0.5,
            format="%.2f",
            help="Statutory holiday hours"
        )
    
    # Generate Button
    st.markdown("---")
    
    if st.button("🚀 Generate Paystub", type="primary"):
        if not uploaded_file:
            st.error("❌ Please upload a DOCX file!")
            return
        
        if not start_date or not end_date:
            st.error("❌ Please select both start and end dates!")
            return
        
        st.session_state.logs = []
        st.session_state.output_docx = None
        
        with st.spinner("🔄 Processing paystub with optimized AI..."):
            try:
                def add_log(msg):
                    st.session_state.logs.append(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}")
                
                doc_bytes = uploaded_file.getvalue()
                doc = Document(io.BytesIO(doc_bytes))
                
                system = OptimizedSystem(st.session_state.api_key, add_log_callback=add_log)
                
                processed_doc, replaced, calcs = system.generate(
                    doc,
                    yearly_income,
                    period,
                    regular_hours,
                    overtime_hours,
                    vacation_hours,
                    stat_holiday_hours,
                    year
                )
                
                st.session_state.last_calcs = calcs
                
                output_buffer = io.BytesIO()
                processed_doc.save(output_buffer)
                output_buffer.seek(0)
                
                st.session_state.output_docx = output_buffer.getvalue()
                st.session_state.output_filename = f"paystub_p{period}_{year}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                
                add_log("")
                add_log(f"💾 Output: {st.session_state.output_filename}")
                add_log(f"✅ Total replacements: {replaced}")
                add_log(f"✅ Period {period} ({year}) calculations applied")
                
                st.success("✅ Paystub generated successfully!")
                st.balloons()
                
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
                add_log(f"❌ ERROR: {str(e)}")
                import traceback
                add_log(f"❌ TRACEBACK: {traceback.format_exc()}")
    
    # Show YTD limit warnings if applicable
    if st.session_state.last_calcs:
        calcs = st.session_state.last_calcs
        if calcs.get('cpp_limit_hit') or calcs.get('ei_limit_hit'):
            st.markdown("---")
            st.markdown('<div class="limit-warning">', unsafe_allow_html=True)
            st.markdown("### ⚠️ YTD Limits Reached")
            if calcs['cpp_limit_hit']:
                st.markdown(f"**CPP**: Maximum YTD limit of ${calcs['cpp_max']:,.2f} has been reached")
            if calcs['ei_limit_hit']:
                st.markdown(f"**EI**: Maximum YTD limit of ${calcs['ei_max']:,.2f} has been reached")
            st.markdown("No further deductions for these items will be made in subsequent pay periods.")
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Display Logs
    if st.session_state.logs:
        st.markdown("---")
        st.markdown("### 📝 Processing Log")
        log_html = "<br>".join([f'<span style="color: #00ff00;">{log}</span>' for log in st.session_state.logs])
        st.markdown(
            f'<div class="log-container">{log_html}</div>',
            unsafe_allow_html=True
        )
    
    # Download Button
    if st.session_state.output_docx:
        st.markdown("---")
        st.markdown("### 📥 Download Processed Paystub")
        
        col_dl1, col_dl2, col_dl3 = st.columns([2, 1, 2])
        
        with col_dl1:
            st.download_button(
                label="⬇️ Download Generated Paystub (DOCX)",
                data=st.session_state.output_docx,
                file_name=st.session_state.output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        
        with col_dl2:
            file_size = len(st.session_state.output_docx) / 1024
            st.metric("File Size", f"{file_size:.2f} KB")
        
        with col_dl3:
            if st.session_state.last_calcs:
                st.metric("Period", f"{st.session_state.last_calcs['period']} of 26")
    
    # Show calculation summary
    if st.session_state.last_calcs:
        st.markdown("---")
        st.markdown("### 📊 Calculation Summary")
        
        calcs = st.session_state.last_calcs
        
        col_sum1, col_sum2, col_sum3, col_sum4 = st.columns(4)
        
        with col_sum1:
            st.metric("Gross Pay", f"${calcs['gross_current']:,.2f}")
            st.caption(f"YTD: ${calcs['gross_ytd']:,.2f}")
        
        with col_sum2:
            st.metric("Net Pay", f"${calcs['net_current']:,.2f}")
            st.caption(f"YTD: ${calcs['net_ytd']:,.2f}")
        
        with col_sum3:
            cpp_label = "CPP (5.7%)"
            if calcs['cpp_limit_hit']:
                cpp_label += " ⚠️ MAX"
            st.metric(cpp_label, f"${calcs['cpp_current']:,.2f}")
            st.caption(f"YTD: ${calcs['cpp_ytd']:,.2f}")
        
        with col_sum4:
            ei_label = "EI (1.7%)"
            if calcs['ei_limit_hit']:
                ei_label += " ⚠️ MAX"
            st.metric(ei_label, f"${calcs['ei_current']:,.2f}")
            st.caption(f"YTD: ${calcs['ei_ytd']:,.2f}")
        
        with st.expander("📋 View Detailed Breakdown"):
            st.markdown("#### Current Period")
            st.write(f"- **Hourly Rate**: ${calcs['hourly_rate']:.2f}")
            st.write(f"- **Regular Hours**: {calcs['hours']:.1f} hrs × ${calcs['hourly_rate']:.2f} = ${calcs['regular_current']:,.2f}")
            if calcs['overtime_hours'] > 0:
                st.write(f"- **Overtime Hours**: {calcs['overtime_hours']:.1f} hrs × ${calcs['overtime_rate']:.2f} = ${calcs['overtime_current']:,.2f}")
            if calcs['vacation_hours'] > 0:
                st.write(f"- **Vacation Hours**: {calcs['vacation_hours']:.1f} hrs")
            if calcs['stat_holidays'] > 0:
                st.write(f"- **Stat Holiday**: {calcs['stat_holidays']:.1f} hrs × ${calcs['hourly_rate']:.2f} = ${calcs['stat_holiday_current']:,.2f}")
            st.write(f"- **Vacation Pay (1.25%)**: ${calcs['vacation_current']:.2f}")
            st.write(f"- **Gross Pay**: ${calcs['gross_current']:,.2f}")
            
            st.markdown("#### Deductions")
            st.write(f"- **CPP (5.7%)**: ${calcs['cpp_current']:,.2f}")
            st.write(f"- **EI (1.7%)**: ${calcs['ei_current']:,.2f}")
            st.write(f"- **Federal Tax (15%)**: ${calcs['tax_current']:,.2f}")
            st.write(f"- **Total Deductions**: ${calcs['deductions_current']:,.2f}")
            
            st.markdown("#### Year to Date")
            st.write(f"- **Gross YTD**: ${calcs['gross_ytd']:,.2f}")
            st.write(f"- **CPP YTD**: ${calcs['cpp_ytd']:,.2f} (Max: ${calcs['cpp_max']:,.2f})")
            st.write(f"- **EI YTD**: ${calcs['ei_ytd']:,.2f} (Max: ${calcs['ei_max']:,.2f})")
            st.write(f"- **Tax YTD**: ${calcs['tax_ytd']:,.2f}")
            st.write(f"- **Net YTD**: ${calcs['net_ytd']:,.2f}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; font-size: 0.9rem;">
        <p><strong>Paystub Generator v5.3 FIXED</strong></p>
        <p>✅ Single AI Call • Safe Document Handling • No Corruption</p>
        <p>✅ CPP 5.7% • EI 1.7% • Vacation 1.25% • Tax 15%</p>
        <p>⚠️ YTD Limits: CPP & EI max out around period 17-18</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()









